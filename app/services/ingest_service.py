import io
import fitz  # PyMuPDF
import docx
import pandas as pd
import pytesseract
from PIL import Image
from fastapi import APIRouter, UploadFile, File, Form
from pydantic import BaseModel
import chromadb

router = APIRouter(prefix="/ingest", tags=["Ingest"])

# Persistent ChromaDB client
chroma_client = chromadb.PersistentClient(path="./chroma_db")

# Create/get collection
collection = chroma_client.get_or_create_collection("kai_collection")

# --------------------------
# Chunking helper
# --------------------------
def chunk_text(text: str, chunk_size: int = 500):
    """Split text into chunks of approx. chunk_size words."""
    words = text.split()
    chunks = [" ".join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)]
    return chunks

# --------------------------
# Text ingestion model
# --------------------------
class Document(BaseModel):
    doc_id: str
    text: str

@router.post("/")
async def ingest_document(doc: Document):
    """Ingest plain text into the vector DB (with chunking)."""
    try:
        chunks = chunk_text(doc.text)
        ids = [f"{doc.doc_id}_{i}" for i in range(len(chunks))]
        collection.add(documents=chunks, ids=ids)
        return {"message": f"Document {doc.doc_id} ingested successfully!", "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}

# --------------------------
# File ingestion
# --------------------------
@router.post("/file")
async def ingest_file(doc_id: str = Form(...), file: UploadFile = File(...)):
    """
    Ingest PDFs, DOCX, CSVs, or images.
    Extract text, chunk it, and push into ChromaDB.
    """
    try:
        filename = file.filename.lower()
        contents = await file.read()
        text = ""

        if filename.endswith(".pdf"):
            pdf = fitz.open(stream=contents, filetype="pdf")
            text = "\n".join(page.get_text() for page in pdf)

        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(contents))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif filename.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(contents))
            text = df.to_string(index=False)

        elif filename.endswith((".png", ".jpg", ".jpeg")):
            image = Image.open(io.BytesIO(contents))
            text = pytesseract.image_to_string(image)

        else:
            return {"error": "Unsupported file format."}

        if not text.strip():
            return {"error": "No text extracted from file."}

        # Chunk text and add to ChromaDB
        chunks = chunk_text(text)
        ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        collection.add(documents=chunks, ids=ids)

        return {"message": f"File {file.filename} ingested successfully!", "chunks": len(chunks), "chars": len(text)}

    except Exception as e:
        return {"error": str(e)}
